import json
import tkinter as tk
from tkinter import filedialog, messagebox, colorchooser, ttk
from pathlib import Path
import os

CONFIG_FILE = "config.json"

try:
    from docx import Document
    from docx.shared import RGBColor, Cm
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import win32com.client
    WIN32COM_AVAILABLE = True
except ImportError:
    WIN32COM_AVAILABLE = False


class ChatExtractorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("对话记录提取工具 → TXT / Word")
        self.root.geometry("880x750")
        self.root.resizable(False, False)

        # ---------- 变量 ----------
        self.input_path = tk.StringVar()
        self.txt_output_path = tk.StringVar()
        self.docx_output_path = tk.StringVar()
        self.role1 = tk.StringVar(value="我")
        self.role2 = tk.StringVar(value="你")

        self.enable_txt = tk.BooleanVar(value=True)
        self.enable_docx = tk.BooleanVar(value=False)

        self.gap_within = tk.BooleanVar(value=False)
        self.gap_between = tk.BooleanVar(value=True)

        self.use_background = tk.BooleanVar(value=False)
        self.bg_image_path = tk.StringVar()
        self.font_color = tk.StringVar(value="black")

        # 目录记忆
        self.last_input_dir = tk.StringVar(value="")
        self.last_bg_dir = tk.StringVar(value="")

        self.load_config()
        self.create_widgets()

    def create_widgets(self):
        row = 0

        # 输入 JSON 文件
        tk.Label(self.root, text="输入 JSON 文件：").grid(row=row, column=0, padx=10, pady=10, sticky='e')
        tk.Entry(self.root, textvariable=self.input_path, width=50).grid(row=row, column=1, padx=5, pady=10)
        tk.Button(self.root, text="浏览...", command=self.select_input).grid(row=row, column=2, padx=5, pady=10)
        row += 1

        # 输出格式
        tk.Label(self.root, text="输出格式：").grid(row=row, column=0, padx=10, pady=10, sticky='ne')
        fmt_frame = tk.Frame(self.root)
        fmt_frame.grid(row=row, column=1, sticky='w', padx=5, pady=10)
        tk.Checkbutton(fmt_frame, text="TXT", variable=self.enable_txt).pack(side=tk.LEFT, padx=5)
        tk.Checkbutton(fmt_frame, text="Word (.docx)", variable=self.enable_docx).pack(side=tk.LEFT, padx=5)
        row += 1

        # TXT 输出路径
        tk.Label(self.root, text="TXT 输出路径：").grid(row=row, column=0, padx=10, pady=5, sticky='e')
        tk.Entry(self.root, textvariable=self.txt_output_path, width=50).grid(row=row, column=1, padx=5, pady=5)
        tk.Button(self.root, text="浏览...", command=lambda: self.select_output('txt')).grid(row=row, column=2, padx=5, pady=5)
        row += 1

        # Word 输出路径
        tk.Label(self.root, text="Word 输出路径：").grid(row=row, column=0, padx=10, pady=5, sticky='e')
        tk.Entry(self.root, textvariable=self.docx_output_path, width=50).grid(row=row, column=1, padx=5, pady=5)
        tk.Button(self.root, text="浏览...", command=lambda: self.select_output('docx')).grid(row=row, column=2, padx=5, pady=5)
        row += 1

        # 角色1
        tk.Label(self.root, text="角色1（prompt）名称：").grid(row=row, column=0, padx=10, pady=10, sticky='e')
        tk.Entry(self.root, textvariable=self.role1, width=20).grid(row=row, column=1, sticky='w', padx=5, pady=10)
        row += 1

        # 角色2
        tk.Label(self.root, text="角色2（expected_output）名称：").grid(row=row, column=0, padx=10, pady=10, sticky='e')
        tk.Entry(self.root, textvariable=self.role2, width=20).grid(row=row, column=1, sticky='w', padx=5, pady=10)
        row += 1

        # 空行控制
        tk.Label(self.root, text="空行控制：").grid(row=row, column=0, padx=10, pady=10, sticky='ne')
        gap_frame = tk.Frame(self.root)
        gap_frame.grid(row=row, column=1, sticky='w', padx=5, pady=10)
        tk.Checkbutton(gap_frame, text="同组内对话间插入空行", variable=self.gap_within).pack(anchor='w')
        tk.Checkbutton(gap_frame, text="各组对话间插入空行", variable=self.gap_between).pack(anchor='w')
        row += 1

        # 字体颜色
        tk.Label(self.root, text="Word 字体颜色：").grid(row=row, column=0, padx=10, pady=10, sticky='e')
        color_frame = tk.Frame(self.root)
        color_frame.grid(row=row, column=1, sticky='w', padx=5, pady=10)
        colors = ['black', 'red', 'blue', 'green', 'purple', 'orange', 'gray', 'darkblue', 'darkred', 'darkgreen']
        self.color_combo = ttk.Combobox(color_frame, textvariable=self.font_color, values=colors, width=12, state='readonly')
        self.color_combo.pack(side=tk.LEFT, padx=5)
        tk.Button(color_frame, text="更多颜色...", command=self.choose_custom_color).pack(side=tk.LEFT, padx=5)
        row += 1

        # 背景图片
        bg_frame = tk.Frame(self.root)
        bg_frame.grid(row=row, column=0, columnspan=3, sticky='w', padx=10, pady=5)
        tk.Checkbutton(bg_frame, text="为Word文档添加背景图片（覆盖整页）", variable=self.use_background).pack(anchor='w')
        tk.Entry(bg_frame, textvariable=self.bg_image_path, width=40).pack(side=tk.LEFT, padx=(20,5))
        tk.Button(bg_frame, text="选择图片...", command=self.select_bg_image).pack(side=tk.LEFT)
        row += 1

        # 转换按钮
        self.convert_btn = tk.Button(self.root, text="开始提取", command=self.convert,
                                     bg="#4CAF50", fg="white", font=("Arial", 12))
        self.convert_btn.grid(row=row, column=1, pady=20)
        row += 1

        # 状态信息框
        self.status_frame = tk.Frame(self.root)
        self.status_frame.grid(row=row, column=0, columnspan=3, padx=10, pady=10, sticky='nsew')
        self.status_text = tk.Text(self.status_frame, height=12, wrap=tk.WORD)
        self.status_scroll = tk.Scrollbar(self.status_frame, orient=tk.VERTICAL, command=self.status_text.yview)
        self.status_text.configure(yscrollcommand=self.status_scroll.set)
        self.status_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.status_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        self.root.grid_rowconfigure(row, weight=1)
        self.root.grid_columnconfigure(1, weight=1)

    # ---------- 回调函数 ----------
    def choose_custom_color(self):
        color = colorchooser.askcolor(title="选择字体颜色")[1]
        if color:
            self.font_color.set(color)

    def select_input(self):
        initial_dir = self.last_input_dir.get() or str(Path.home())
        path = filedialog.askopenfilename(
            title="选择输入 JSON 文件",
            initialdir=initial_dir,
            filetypes=[("JSON files", "*.json")]
        )
        if path:
            self.input_path.set(path)
            self.last_input_dir.set(str(Path(path).parent))
            base = Path(path).with_suffix('')
            self.txt_output_path.set(str(base) + "_converted.txt")
            self.docx_output_path.set(str(base) + "_converted.docx")
            self.save_config()

    def select_output(self, fmt):
        initial_dir = self.last_input_dir.get() or str(Path.home())
        if fmt == 'txt':
            path = filedialog.asksaveasfilename(
                title="保存 TXT 文件",
                initialdir=initial_dir,
                defaultextension=".txt",
                filetypes=[("TXT files", "*.txt")]
            )
            if path:
                self.txt_output_path.set(path)
                self.last_input_dir.set(str(Path(path).parent))
                self.save_config()
        else:
            path = filedialog.asksaveasfilename(
                title="保存 Word 文件",
                initialdir=initial_dir,
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx")]
            )
            if path:
                self.docx_output_path.set(path)
                self.last_input_dir.set(str(Path(path).parent))
                self.save_config()

    def select_bg_image(self):
        initial_dir = self.last_bg_dir.get() or str(Path.home())
        path = filedialog.askopenfilename(
            title="选择背景图片",
            initialdir=initial_dir,
            filetypes=[("Image files", "*.png *.jpg *.jpeg *.bmp *.gif")]
        )
        if path:
            self.bg_image_path.set(path)
            self.last_bg_dir.set(str(Path(path).parent))
            self.save_config()

    def log(self, msg):
        self.status_text.insert(tk.END, msg + "\n")
        self.status_text.see(tk.END)
        self.root.update()

    # ---------- 配置读写 ----------
    def load_config(self):
        if Path(CONFIG_FILE).exists():
            try:
                with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                    cfg = json.load(f)
                self.role1.set(cfg.get("role1", "我"))
                self.role2.set(cfg.get("role2", "你"))
                self.gap_within.set(cfg.get("gap_within", False))
                self.gap_between.set(cfg.get("gap_between", True))
                self.font_color.set(cfg.get("font_color", "black"))
                self.use_background.set(cfg.get("use_background", False))
                self.bg_image_path.set(cfg.get("bg_image_path", ""))
                self.enable_txt.set(cfg.get("enable_txt", True))
                self.enable_docx.set(cfg.get("enable_docx", False))
                self.txt_output_path.set(cfg.get("txt_output_path", ""))
                self.docx_output_path.set(cfg.get("docx_output_path", ""))
                self.last_input_dir.set(cfg.get("last_input_dir", ""))
                self.last_bg_dir.set(cfg.get("last_bg_dir", ""))
            except:
                pass

    def save_config(self):
        cfg = {
            "role1": self.role1.get(),
            "role2": self.role2.get(),
            "gap_within": self.gap_within.get(),
            "gap_between": self.gap_between.get(),
            "font_color": self.font_color.get(),
            "use_background": self.use_background.get(),
            "bg_image_path": self.bg_image_path.get(),
            "enable_txt": self.enable_txt.get(),
            "enable_docx": self.enable_docx.get(),
            "txt_output_path": self.txt_output_path.get(),
            "docx_output_path": self.docx_output_path.get(),
            "last_input_dir": self.last_input_dir.get(),
            "last_bg_dir": self.last_bg_dir.get()
        }
        with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump(cfg, f, ensure_ascii=False, indent=2)

    # ---------- 核心转换 ----------
    def convert(self):
        # 校验
        input_path = self.input_path.get().strip()
        if not input_path:
            messagebox.showerror("错误", "请选择输入 JSON 文件")
            return
        if not self.enable_txt.get() and not self.enable_docx.get():
            messagebox.showerror("错误", "请至少选择一种输出格式")
            return
        if self.enable_txt.get() and not self.txt_output_path.get().strip():
            messagebox.showerror("错误", "请指定 TXT 输出路径")
            return
        if self.enable_docx.get() and not self.docx_output_path.get().strip():
            messagebox.showerror("错误", "请指定 Word 输出路径")
            return
        if self.enable_docx.get() and not DOCX_AVAILABLE:
            messagebox.showerror("错误", "生成 Word 需要 python-docx。\npip install python-docx\n安装后请重启程序并选择正确的解释器")
            return
        if self.use_background.get() and not self.bg_image_path.get().strip():
            messagebox.showerror("错误", "已勾选背景图片，但未选择图片。")
            return
        if self.use_background.get() and not self.enable_docx.get():
            messagebox.showwarning("提示", "背景仅适用于 Word，您未勾选，背景将被忽略。")
        if self.use_background.get() and not WIN32COM_AVAILABLE:
            messagebox.showerror("错误", "添加背景需要 pywin32。\npip install pywin32\n安装后请重启程序并选择正确的解释器")
            return

        # 读取 JSON
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
        except Exception as e:
            messagebox.showerror("错误", f"读取 JSON 失败：{e}")
            return
        if 'evals' not in data:
            messagebox.showerror("错误", "JSON 中缺少 'evals' 字段")
            return

        # 构建对话对
        pairs = []
        for item in data['evals']:
            prompt = item.get('prompt', '').strip()
            expected = item.get('expected_output', '').strip()
            pair = []
            if prompt:
                pair.append(f"{self.role1.get()}：{prompt}")
            if expected:
                pair.append(f"{self.role2.get()}：{expected}")
            if pair:
                pairs.append(pair)

        if not pairs:
            messagebox.showwarning("警告", "未提取到任何有效对话记录")
            return

        self.convert_btn.config(state=tk.DISABLED)
        self.log("开始处理...")

        try:
            # ---- TXT ----
            if self.enable_txt.get():
                txt_path = self.txt_output_path.get().strip()
                with open(txt_path, 'w', encoding='utf-8') as f:
                    for idx, pair in enumerate(pairs):
                        if self.gap_within.get() and len(pair) > 1:
                            f.write(pair[0] + "\n\n" + pair[1])
                        else:
                            f.write("\n".join(pair))
                        if idx < len(pairs) - 1:
                            f.write("\n\n" if self.gap_between.get() else "\n")
                self.log(f"✅ TXT 导出完成：{txt_path}")

            # ---- Word ----
            if self.enable_docx.get():
                docx_path = self.docx_output_path.get().strip()
                doc = Document()

                # 字体颜色
                try:
                    color_str = self.font_color.get()
                    if color_str.startswith('#'):
                        r, g, b = int(color_str[1:3], 16), int(color_str[3:5], 16), int(color_str[5:7], 16)
                        rgb = RGBColor(r, g, b)
                    else:
                        color_map = {
                            'black': RGBColor(0,0,0), 'red': RGBColor(255,0,0),
                            'blue': RGBColor(0,0,255), 'green': RGBColor(0,128,0),
                            'purple': RGBColor(128,0,128), 'orange': RGBColor(255,165,0),
                            'gray': RGBColor(128,128,128), 'darkblue': RGBColor(0,0,139),
                            'darkred': RGBColor(139,0,0), 'darkgreen': RGBColor(0,100,0)
                        }
                        rgb = color_map.get(color_str, RGBColor(0,0,0))
                except:
                    rgb = RGBColor(0,0,0)

                # 缩进值（启用背景时左右均缩进 2.54cm）
                indent = Cm(2.54) if self.use_background.get() else 0

                for idx, pair in enumerate(pairs):
                    for line in pair:
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = indent
                        p.paragraph_format.right_indent = indent
                        run = p.add_run(line)
                        run.font.color.rgb = rgb
                    if idx < len(pairs) - 1 and self.gap_between.get():
                        doc.add_paragraph()

                # 先保存临时文档
                temp_docx = docx_path + ".tmp.docx"
                doc.save(temp_docx)

                # ---- 使用 win32com 添加背景图片 ----
                if self.use_background.get() and self.bg_image_path.get().strip():
                    try:
                        self.add_background_with_word(temp_docx, docx_path, self.bg_image_path.get().strip())
                        os.remove(temp_docx)
                        self.log(f"✅ Word 导出完成（含背景图片）：{docx_path}")
                    except Exception as e:
                        self.log(f"⚠️ 添加背景失败，保留无背景版本：{e}")
                        os.replace(temp_docx, docx_path)
                        self.log(f"✅ Word 导出完成（无背景）：{docx_path}")
                else:
                    os.replace(temp_docx, docx_path)
                    self.log(f"✅ Word 导出完成：{docx_path}")

            self.save_config()
            messagebox.showinfo("完成", "所有导出任务已完成！")
        except Exception as e:
            self.log(f"❌ 错误：{str(e)}")
            messagebox.showerror("错误", f"处理失败：{str(e)}")
        finally:
            self.convert_btn.config(state=tk.NORMAL)

    # ---------- 使用 win32com 添加背景图片（溢出覆盖法） ----------
    def add_background_with_word(self, input_path, output_path, bg_image_path):
        """
        使用 WPS COM 在每一页的页眉中插入浮动图片，
        将图片尺寸设得比页面稍大，并略微偏移，使其四周溢出，完全覆盖页面。
        """
        word = None
        doc = None
        try:
            word = win32com.client.Dispatch("Kwps.Application")
            word.Visible = False
            doc = word.Documents.Open(os.path.abspath(input_path))

            # 遍历所有节
            for section in doc.Sections:
                # 设置页面边距为0
                section.PageSetup.LeftMargin = 0
                section.PageSetup.RightMargin = 0
                section.PageSetup.TopMargin = 0
                section.PageSetup.BottomMargin = 0
                try:
                    section.PageSetup.HeaderDistance = 0
                except:
                    pass

                # 获取页眉
                header = section.Headers(1)
                header.Range.Delete()

                # 获取页面尺寸（单位：磅）
                page_width = section.PageSetup.PageWidth
                page_height = section.PageSetup.PageHeight

                # 计算溢出量（例如 20 磅）
                overflow = 20
                img_width = page_width + overflow
                img_height = page_height + overflow
                # 偏移量为负溢出量的一半，使图片居中溢出
                left_offset = -overflow / 2
                top_offset = -overflow / 2

                # 添加浮动图片
                shape = header.Shapes.AddPicture(
                    FileName=os.path.abspath(bg_image_path),
                    LinkToFile=False,
                    SaveWithDocument=True,
                    Left=left_offset,
                    Top=top_offset,
                    Width=img_width,
                    Height=img_height,
                    Anchor=header.Range
                )

                # 设置环绕方式：衬于文字下方
                shape.WrapFormat.Type = 3  # wdWrapBehind

                # 设置相对位置：相对于页面
                shape.RelativeHorizontalPosition = 0
                shape.RelativeVerticalPosition = 0

                # 确保位置为计算值
                shape.Left = left_offset
                shape.Top = top_offset

                # 解除锁定纵横比，使图片严格拉伸
                shape.LockAspectRatio = False

            # 保存
            doc.SaveAs(os.path.abspath(output_path))
            doc.Close()
        except Exception as e:
            if doc:
                doc.Close()
            raise e
        finally:
            if word:
                word.Quit()


if __name__ == "__main__":
    root = tk.Tk()
    app = ChatExtractorApp(root)
    root.mainloop()